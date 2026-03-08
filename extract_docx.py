from docx import Document
import json

doc = Document(r'S:\Iniyanmozhi\iniyenmozhi-dreams\Career Restart Program for Young Women.docx')

content = []
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        content.append({
            "index": i,
            "text": para.text.strip(),
            "style": para.style.name if para.style else "Normal"
        })

# Also check for tables (forms)
tables_data = []
for t_idx, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    tables_data.append({"table_index": t_idx, "rows": rows})

result = {"paragraphs": content, "tables": tables_data}

with open(r'S:\Iniyanmozhi\iniyenmozhi-dreams\docx_content.json', 'w', encoding='utf-8') as f:
    json.dump(result, f, indent=2, ensure_ascii=False)

print(f"Done! {len(content)} paragraphs, {len(tables_data)} tables")
